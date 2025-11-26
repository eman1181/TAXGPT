from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import csv
from datetime import datetime
import os


from GRAD_FINAL_BACKEND import generate_final_tax_reply

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class TaxRequest(BaseModel):
    question: str

class ExportRequest(BaseModel):
    content: str  # نص الخطاب بعد التعديل

class FeedbackRequest(BaseModel):
    question: str
    original_answer: str          # رد الموديل زي ما طلع
    edited_answer: str | None = None  # الرد بعد تعديل المستخدم (لو عدّل)
    rating: int                   # 1 = مفيد ، 0 = غير دقيق
    comment: str | None = None



@app.post("/tax-answer")
def tax_answer(body: TaxRequest):
    answer = generate_final_tax_reply(body.question)
    return {"answer": answer}

# ✅ ده اللي بيحوّل النص إلى ملف Word
@app.post("/export-docx")
def export_docx(body: ExportRequest):
    doc = Document()

    # نمرّ على كل سطر ونضيفه كفقرة RTL ومحاذاة لليمين
    for line in body.content.split("\n"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # محاذاة لليمين

        run = p.add_run(line)

        # تفعيل اتجاه من اليمين لليسار (RTL) في الـ run
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement("w:rtl")
        rPr.append(rtl)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type=(
            "application/vnd.openxmlformats-"
            "officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": 'attachment; filename="tax_letter.docx"'},
    )

FEEDBACK_FILE = "feedback_with_edits.csv"


@app.post("/feedback")
def save_feedback(body: FeedbackRequest):
    file_exists = os.path.isfile(FEEDBACK_FILE)

    with open(FEEDBACK_FILE, mode="a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if not file_exists:
            writer.writerow([
                "timestamp",
                "question",
                "original_answer",
                "edited_answer",
                "rating",
                "comment",
            ])

        writer.writerow([
            datetime.now().isoformat(),
            body.question,
            body.original_answer,
            body.edited_answer or "",
            body.rating,
            body.comment or "",
        ])

    return {"status": "ok"}

@app.get("/ping")
def ping():
    return {"status": "ok"}
